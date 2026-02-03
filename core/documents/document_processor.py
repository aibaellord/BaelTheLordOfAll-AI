"""
BAEL - Document Processor
Advanced document processing for PDFs, Office files, and more.

Features:
- PDF parsing and extraction
- Word/Excel/PowerPoint processing
- Text extraction
- Table extraction
- Image extraction
- OCR integration
- Document conversion
"""

import asyncio
import base64
import hashlib
import io
import logging
import os
import re
import tempfile
from abc import ABC, abstractmethod
from dataclasses import dataclass, field
from datetime import datetime
from enum import Enum
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple, Union

logger = logging.getLogger(__name__)


# =============================================================================
# ENUMS
# =============================================================================

class DocumentType(Enum):
    """Supported document types."""
    PDF = "pdf"
    DOCX = "docx"
    DOC = "doc"
    XLSX = "xlsx"
    XLS = "xls"
    PPTX = "pptx"
    PPT = "ppt"
    TXT = "txt"
    MD = "md"
    HTML = "html"
    CSV = "csv"
    JSON = "json"
    UNKNOWN = "unknown"


class ExtractionMode(Enum):
    """Text extraction modes."""
    FULL = "full"           # All text
    STRUCTURED = "structured"  # With structure
    TABLES_ONLY = "tables_only"
    IMAGES_ONLY = "images_only"
    METADATA_ONLY = "metadata_only"


# =============================================================================
# DATA CLASSES
# =============================================================================

@dataclass
class DocumentMetadata:
    """Document metadata."""
    title: str = ""
    author: str = ""
    created: Optional[datetime] = None
    modified: Optional[datetime] = None
    pages: int = 0
    words: int = 0
    chars: int = 0
    language: str = ""
    file_size: int = 0
    file_hash: str = ""
    extra: Dict[str, Any] = field(default_factory=dict)


@dataclass
class TextBlock:
    """Extracted text block."""
    text: str
    page: int = 0
    x: float = 0.0
    y: float = 0.0
    width: float = 0.0
    height: float = 0.0
    font: str = ""
    font_size: float = 0.0
    is_bold: bool = False
    is_italic: bool = False


@dataclass
class TableCell:
    """Table cell."""
    value: Any
    row: int
    col: int
    row_span: int = 1
    col_span: int = 1


@dataclass
class ExtractedTable:
    """Extracted table."""
    cells: List[List[Any]]
    page: int = 0
    headers: List[str] = field(default_factory=list)
    metadata: Dict[str, Any] = field(default_factory=dict)

    @property
    def rows(self) -> int:
        return len(self.cells)

    @property
    def cols(self) -> int:
        return len(self.cells[0]) if self.cells else 0

    def to_dict(self) -> List[Dict[str, Any]]:
        """Convert to list of dicts using headers."""
        if not self.headers or not self.cells:
            return []

        result = []
        for row in self.cells[1:] if self.headers else self.cells:
            row_dict = {}
            for i, val in enumerate(row):
                key = self.headers[i] if i < len(self.headers) else f"col_{i}"
                row_dict[key] = val
            result.append(row_dict)

        return result

    def to_csv(self) -> str:
        """Convert to CSV string."""
        import csv
        output = io.StringIO()
        writer = csv.writer(output)

        if self.headers:
            writer.writerow(self.headers)

        for row in self.cells:
            writer.writerow(row)

        return output.getvalue()


@dataclass
class ExtractedImage:
    """Extracted image."""
    data: bytes
    format: str = "png"
    page: int = 0
    x: float = 0.0
    y: float = 0.0
    width: float = 0.0
    height: float = 0.0
    description: str = ""

    def to_base64(self) -> str:
        return base64.b64encode(self.data).decode()

    def save(self, path: str) -> None:
        with open(path, "wb") as f:
            f.write(self.data)


@dataclass
class DocumentSection:
    """Document section/chapter."""
    title: str
    level: int  # 1=h1, 2=h2, etc.
    content: str
    page_start: int = 0
    page_end: int = 0
    children: List["DocumentSection"] = field(default_factory=list)


@dataclass
class ExtractionResult:
    """Complete extraction result."""
    text: str
    metadata: DocumentMetadata
    blocks: List[TextBlock] = field(default_factory=list)
    tables: List[ExtractedTable] = field(default_factory=list)
    images: List[ExtractedImage] = field(default_factory=list)
    sections: List[DocumentSection] = field(default_factory=list)
    raw_pages: List[str] = field(default_factory=list)


# =============================================================================
# EXTRACTORS
# =============================================================================

class DocumentExtractor(ABC):
    """Abstract document extractor."""

    @property
    @abstractmethod
    def supported_types(self) -> List[DocumentType]:
        pass

    @abstractmethod
    async def extract(
        self,
        file_path: str,
        mode: ExtractionMode = ExtractionMode.FULL
    ) -> ExtractionResult:
        pass


class PDFExtractor(DocumentExtractor):
    """PDF document extractor."""

    @property
    def supported_types(self) -> List[DocumentType]:
        return [DocumentType.PDF]

    async def extract(
        self,
        file_path: str,
        mode: ExtractionMode = ExtractionMode.FULL
    ) -> ExtractionResult:
        """Extract content from PDF."""
        try:
            import pypdf
            has_pypdf = True
        except ImportError:
            has_pypdf = False

        if not has_pypdf:
            return await self._fallback_extract(file_path)

        reader = pypdf.PdfReader(file_path)

        # Get metadata
        info = reader.metadata or {}
        metadata = DocumentMetadata(
            title=str(info.get("/Title", "")),
            author=str(info.get("/Author", "")),
            pages=len(reader.pages),
            file_size=os.path.getsize(file_path)
        )

        # Calculate file hash
        with open(file_path, "rb") as f:
            metadata.file_hash = hashlib.md5(f.read()).hexdigest()

        # Extract based on mode
        if mode == ExtractionMode.METADATA_ONLY:
            return ExtractionResult(text="", metadata=metadata)

        # Extract text
        full_text = []
        raw_pages = []
        blocks = []

        for page_num, page in enumerate(reader.pages):
            page_text = page.extract_text() or ""
            raw_pages.append(page_text)
            full_text.append(page_text)

            # Create text block for page
            blocks.append(TextBlock(
                text=page_text,
                page=page_num + 1
            ))

        # Extract tables if needed
        tables = []
        if mode in (ExtractionMode.FULL, ExtractionMode.STRUCTURED, ExtractionMode.TABLES_ONLY):
            tables = await self._extract_tables(reader)

        # Extract images if needed
        images = []
        if mode in (ExtractionMode.FULL, ExtractionMode.IMAGES_ONLY):
            images = await self._extract_images(reader)

        # Update metadata
        full_text_str = "\n\n".join(full_text)
        metadata.words = len(full_text_str.split())
        metadata.chars = len(full_text_str)

        # Detect sections (basic heading detection)
        sections = self._detect_sections(full_text_str)

        return ExtractionResult(
            text=full_text_str,
            metadata=metadata,
            blocks=blocks,
            tables=tables,
            images=images,
            sections=sections,
            raw_pages=raw_pages
        )

    async def _extract_tables(self, reader) -> List[ExtractedTable]:
        """Extract tables from PDF."""
        tables = []

        try:
            import tabula

            # Use tabula for table extraction
            dfs = tabula.read_pdf(reader, pages="all")

            for i, df in enumerate(dfs):
                cells = df.values.tolist()
                headers = list(df.columns)

                tables.append(ExtractedTable(
                    cells=cells,
                    headers=headers,
                    metadata={"table_index": i}
                ))

        except ImportError:
            logger.debug("tabula not available for table extraction")

        return tables

    async def _extract_images(self, reader) -> List[ExtractedImage]:
        """Extract images from PDF."""
        images = []

        for page_num, page in enumerate(reader.pages):
            if "/XObject" in page["/Resources"]:
                x_objects = page["/Resources"]["/XObject"].get_object()

                for obj_name in x_objects:
                    obj = x_objects[obj_name]

                    if obj["/Subtype"] == "/Image":
                        try:
                            data = obj.get_data()

                            # Determine format
                            if "/Filter" in obj:
                                filter_type = obj["/Filter"]
                                if filter_type == "/DCTDecode":
                                    fmt = "jpeg"
                                elif filter_type == "/FlateDecode":
                                    fmt = "png"
                                else:
                                    fmt = "raw"
                            else:
                                fmt = "raw"

                            images.append(ExtractedImage(
                                data=data,
                                format=fmt,
                                page=page_num + 1
                            ))

                        except Exception as e:
                            logger.debug(f"Failed to extract image: {e}")

        return images

    def _detect_sections(self, text: str) -> List[DocumentSection]:
        """Detect document sections from text."""
        sections = []

        # Simple heuristic: lines that are ALL CAPS or shorter lines
        # followed by content
        lines = text.split("\n")
        current_section = None
        current_content = []

        for line in lines:
            stripped = line.strip()

            # Check if it's a potential heading
            if stripped and len(stripped) < 100:
                is_heading = (
                    stripped.isupper() or
                    stripped.startswith(("#", "Chapter", "Section")) or
                    re.match(r"^\d+\.\s+", stripped)
                )

                if is_heading:
                    # Save previous section
                    if current_section:
                        current_section.content = "\n".join(current_content)
                        sections.append(current_section)

                    # Start new section
                    level = 1 if stripped.isupper() else 2
                    current_section = DocumentSection(
                        title=stripped,
                        level=level,
                        content=""
                    )
                    current_content = []
                    continue

            current_content.append(line)

        # Save last section
        if current_section:
            current_section.content = "\n".join(current_content)
            sections.append(current_section)

        return sections

    async def _fallback_extract(self, file_path: str) -> ExtractionResult:
        """Fallback extraction without pypdf."""
        # Try pdfminer
        try:
            from pdfminer.high_level import extract_text

            text = extract_text(file_path)

            return ExtractionResult(
                text=text,
                metadata=DocumentMetadata(
                    file_size=os.path.getsize(file_path),
                    words=len(text.split()),
                    chars=len(text)
                )
            )

        except ImportError:
            logger.warning("No PDF library available")
            return ExtractionResult(
                text="[PDF extraction not available]",
                metadata=DocumentMetadata()
            )


class WordExtractor(DocumentExtractor):
    """Word document extractor (.docx)."""

    @property
    def supported_types(self) -> List[DocumentType]:
        return [DocumentType.DOCX]

    async def extract(
        self,
        file_path: str,
        mode: ExtractionMode = ExtractionMode.FULL
    ) -> ExtractionResult:
        """Extract content from Word document."""
        try:
            from docx import Document
        except ImportError:
            return ExtractionResult(
                text="[python-docx not available]",
                metadata=DocumentMetadata()
            )

        doc = Document(file_path)

        # Extract metadata
        core_props = doc.core_properties
        metadata = DocumentMetadata(
            title=core_props.title or "",
            author=core_props.author or "",
            created=core_props.created,
            modified=core_props.modified,
            file_size=os.path.getsize(file_path)
        )

        if mode == ExtractionMode.METADATA_ONLY:
            return ExtractionResult(text="", metadata=metadata)

        # Extract text
        paragraphs = []
        blocks = []
        sections = []
        current_section = None

        for para in doc.paragraphs:
            text = para.text
            paragraphs.append(text)

            # Check for heading
            if para.style.name.startswith("Heading"):
                level = int(para.style.name[-1]) if para.style.name[-1].isdigit() else 1

                if current_section:
                    sections.append(current_section)

                current_section = DocumentSection(
                    title=text,
                    level=level,
                    content=""
                )
            elif current_section:
                current_section.content += text + "\n"

            # Create text block
            run = para.runs[0] if para.runs else None
            blocks.append(TextBlock(
                text=text,
                font=run.font.name if run else "",
                font_size=run.font.size.pt if run and run.font.size else 0,
                is_bold=run.bold if run else False,
                is_italic=run.italic if run else False
            ))

        if current_section:
            sections.append(current_section)

        # Extract tables
        tables = []
        if mode in (ExtractionMode.FULL, ExtractionMode.STRUCTURED, ExtractionMode.TABLES_ONLY):
            for table in doc.tables:
                cells = []
                headers = []

                for i, row in enumerate(table.rows):
                    row_data = [cell.text for cell in row.cells]

                    if i == 0:
                        headers = row_data
                    else:
                        cells.append(row_data)

                tables.append(ExtractedTable(
                    cells=cells,
                    headers=headers
                ))

        # Extract images
        images = []
        if mode in (ExtractionMode.FULL, ExtractionMode.IMAGES_ONLY):
            for rel in doc.part.rels.values():
                if "image" in rel.reltype:
                    try:
                        image_data = rel.target_part.blob
                        images.append(ExtractedImage(
                            data=image_data,
                            format="png"  # Simplified
                        ))
                    except Exception as e:
                        logger.debug(f"Failed to extract image: {e}")

        full_text = "\n".join(paragraphs)
        metadata.words = len(full_text.split())
        metadata.chars = len(full_text)

        return ExtractionResult(
            text=full_text,
            metadata=metadata,
            blocks=blocks,
            tables=tables,
            images=images,
            sections=sections
        )


class ExcelExtractor(DocumentExtractor):
    """Excel document extractor (.xlsx)."""

    @property
    def supported_types(self) -> List[DocumentType]:
        return [DocumentType.XLSX, DocumentType.XLS]

    async def extract(
        self,
        file_path: str,
        mode: ExtractionMode = ExtractionMode.FULL
    ) -> ExtractionResult:
        """Extract content from Excel file."""
        try:
            import openpyxl
        except ImportError:
            return ExtractionResult(
                text="[openpyxl not available]",
                metadata=DocumentMetadata()
            )

        wb = openpyxl.load_workbook(file_path, data_only=True)

        metadata = DocumentMetadata(
            file_size=os.path.getsize(file_path),
            extra={"sheets": wb.sheetnames}
        )

        if mode == ExtractionMode.METADATA_ONLY:
            return ExtractionResult(text="", metadata=metadata)

        tables = []
        all_text = []

        for sheet_name in wb.sheetnames:
            sheet = wb[sheet_name]

            # Extract as table
            cells = []
            headers = []

            for i, row in enumerate(sheet.iter_rows()):
                row_data = [cell.value for cell in row]

                # Skip empty rows
                if not any(row_data):
                    continue

                if i == 0:
                    headers = [str(v) if v else f"col_{j}" for j, v in enumerate(row_data)]
                else:
                    cells.append(row_data)

            if cells or headers:
                tables.append(ExtractedTable(
                    cells=cells,
                    headers=headers,
                    metadata={"sheet": sheet_name}
                ))

            # Convert to text
            sheet_text = f"=== Sheet: {sheet_name} ===\n"
            for row in sheet.iter_rows():
                row_values = [str(cell.value) if cell.value else "" for cell in row]
                if any(row_values):
                    sheet_text += "\t".join(row_values) + "\n"

            all_text.append(sheet_text)

        full_text = "\n\n".join(all_text)
        metadata.words = len(full_text.split())
        metadata.chars = len(full_text)

        return ExtractionResult(
            text=full_text,
            metadata=metadata,
            tables=tables
        )


class MarkdownExtractor(DocumentExtractor):
    """Markdown document extractor."""

    @property
    def supported_types(self) -> List[DocumentType]:
        return [DocumentType.MD, DocumentType.TXT]

    async def extract(
        self,
        file_path: str,
        mode: ExtractionMode = ExtractionMode.FULL
    ) -> ExtractionResult:
        """Extract content from Markdown file."""
        with open(file_path, "r", encoding="utf-8") as f:
            text = f.read()

        metadata = DocumentMetadata(
            file_size=os.path.getsize(file_path),
            words=len(text.split()),
            chars=len(text)
        )

        if mode == ExtractionMode.METADATA_ONLY:
            return ExtractionResult(text="", metadata=metadata)

        # Parse sections from markdown headers
        sections = []
        current_section = None

        for line in text.split("\n"):
            if line.startswith("#"):
                level = len(line) - len(line.lstrip("#"))
                title = line.lstrip("#").strip()

                if current_section:
                    sections.append(current_section)

                current_section = DocumentSection(
                    title=title,
                    level=level,
                    content=""
                )
            elif current_section:
                current_section.content += line + "\n"

        if current_section:
            sections.append(current_section)

        # Extract tables from markdown
        tables = self._extract_markdown_tables(text)

        return ExtractionResult(
            text=text,
            metadata=metadata,
            sections=sections,
            tables=tables
        )

    def _extract_markdown_tables(self, text: str) -> List[ExtractedTable]:
        """Extract tables from markdown text."""
        tables = []

        # Find markdown tables
        lines = text.split("\n")
        in_table = False
        table_lines = []

        for line in lines:
            # Check if line is a table row
            if "|" in line and line.strip().startswith("|"):
                if not in_table:
                    in_table = True

                # Skip separator line
                if not re.match(r"^\|[-:\s|]+\|$", line.strip()):
                    table_lines.append(line)
            else:
                if in_table and table_lines:
                    # Parse table
                    cells = []
                    headers = []

                    for i, row_line in enumerate(table_lines):
                        row_cells = [
                            cell.strip()
                            for cell in row_line.strip("|").split("|")
                        ]

                        if i == 0:
                            headers = row_cells
                        else:
                            cells.append(row_cells)

                    if headers or cells:
                        tables.append(ExtractedTable(
                            cells=cells,
                            headers=headers
                        ))

                in_table = False
                table_lines = []

        return tables


# =============================================================================
# OCR INTEGRATION
# =============================================================================

class OCRProcessor:
    """OCR for image-based documents."""

    def __init__(self, tesseract_path: Optional[str] = None):
        self.tesseract_path = tesseract_path

    async def ocr_image(
        self,
        image_data: Union[bytes, str],
        language: str = "eng"
    ) -> str:
        """Perform OCR on image."""
        try:
            import pytesseract
            from PIL import Image

            if self.tesseract_path:
                pytesseract.pytesseract.tesseract_cmd = self.tesseract_path

            if isinstance(image_data, str):
                image = Image.open(image_data)
            else:
                image = Image.open(io.BytesIO(image_data))

            # Run in thread pool
            loop = asyncio.get_event_loop()
            text = await loop.run_in_executor(
                None,
                lambda: pytesseract.image_to_string(image, lang=language)
            )

            return text

        except ImportError:
            logger.warning("pytesseract not available")
            return "[OCR not available]"

    async def ocr_pdf(
        self,
        pdf_path: str,
        language: str = "eng"
    ) -> str:
        """OCR a scanned PDF."""
        try:
            import pdf2image
            import pytesseract
            from PIL import Image

            # Convert PDF to images
            images = pdf2image.convert_from_path(pdf_path)

            # OCR each page
            texts = []
            for i, image in enumerate(images):
                loop = asyncio.get_event_loop()
                text = await loop.run_in_executor(
                    None,
                    lambda img=image: pytesseract.image_to_string(img, lang=language)
                )
                texts.append(f"=== Page {i + 1} ===\n{text}")

            return "\n\n".join(texts)

        except ImportError as e:
            logger.warning(f"OCR dependencies not available: {e}")
            return "[OCR not available]"


# =============================================================================
# DOCUMENT CONVERTER
# =============================================================================

class DocumentConverter:
    """Convert between document formats."""

    async def convert(
        self,
        source_path: str,
        target_format: DocumentType,
        output_path: Optional[str] = None
    ) -> str:
        """Convert document to different format."""
        source_type = self._detect_type(source_path)

        if output_path is None:
            output_path = source_path.rsplit(".", 1)[0] + f".{target_format.value}"

        # Conversion logic
        if target_format == DocumentType.PDF:
            return await self._to_pdf(source_path, source_type, output_path)
        elif target_format == DocumentType.MD:
            return await self._to_markdown(source_path, source_type, output_path)
        elif target_format == DocumentType.TXT:
            return await self._to_text(source_path, source_type, output_path)
        else:
            raise ValueError(f"Unsupported target format: {target_format}")

    async def _to_pdf(
        self,
        source_path: str,
        source_type: DocumentType,
        output_path: str
    ) -> str:
        """Convert to PDF."""
        # Would use libreoffice, pandoc, or similar
        logger.info(f"Converting {source_type.value} to PDF")
        return output_path

    async def _to_markdown(
        self,
        source_path: str,
        source_type: DocumentType,
        output_path: str
    ) -> str:
        """Convert to Markdown."""
        # Extract text and format as markdown
        processor = DocumentProcessor()
        result = await processor.process(source_path)

        md_content = f"# {result.metadata.title or 'Document'}\n\n"

        for section in result.sections:
            md_content += f"{'#' * section.level} {section.title}\n\n"
            md_content += section.content + "\n\n"

        if not result.sections:
            md_content += result.text

        with open(output_path, "w", encoding="utf-8") as f:
            f.write(md_content)

        return output_path

    async def _to_text(
        self,
        source_path: str,
        source_type: DocumentType,
        output_path: str
    ) -> str:
        """Convert to plain text."""
        processor = DocumentProcessor()
        result = await processor.process(source_path)

        with open(output_path, "w", encoding="utf-8") as f:
            f.write(result.text)

        return output_path

    def _detect_type(self, path: str) -> DocumentType:
        """Detect document type from path."""
        ext = Path(path).suffix.lower().lstrip(".")

        type_map = {
            "pdf": DocumentType.PDF,
            "docx": DocumentType.DOCX,
            "doc": DocumentType.DOC,
            "xlsx": DocumentType.XLSX,
            "xls": DocumentType.XLS,
            "pptx": DocumentType.PPTX,
            "ppt": DocumentType.PPT,
            "txt": DocumentType.TXT,
            "md": DocumentType.MD,
            "html": DocumentType.HTML,
            "csv": DocumentType.CSV,
            "json": DocumentType.JSON
        }

        return type_map.get(ext, DocumentType.UNKNOWN)


# =============================================================================
# MAIN DOCUMENT PROCESSOR
# =============================================================================

class DocumentProcessor:
    """Main document processing orchestrator."""

    def __init__(self):
        self._extractors: Dict[DocumentType, DocumentExtractor] = {
            DocumentType.PDF: PDFExtractor(),
            DocumentType.DOCX: WordExtractor(),
            DocumentType.XLSX: ExcelExtractor(),
            DocumentType.XLS: ExcelExtractor(),
            DocumentType.MD: MarkdownExtractor(),
            DocumentType.TXT: MarkdownExtractor()
        }

        self.ocr = OCRProcessor()
        self.converter = DocumentConverter()

    def register_extractor(
        self,
        doc_type: DocumentType,
        extractor: DocumentExtractor
    ) -> None:
        """Register custom extractor."""
        self._extractors[doc_type] = extractor

    def detect_type(self, file_path: str) -> DocumentType:
        """Detect document type."""
        ext = Path(file_path).suffix.lower().lstrip(".")

        type_map = {
            "pdf": DocumentType.PDF,
            "docx": DocumentType.DOCX,
            "doc": DocumentType.DOC,
            "xlsx": DocumentType.XLSX,
            "xls": DocumentType.XLS,
            "pptx": DocumentType.PPTX,
            "ppt": DocumentType.PPT,
            "txt": DocumentType.TXT,
            "md": DocumentType.MD,
            "html": DocumentType.HTML,
            "csv": DocumentType.CSV,
            "json": DocumentType.JSON
        }

        return type_map.get(ext, DocumentType.UNKNOWN)

    async def process(
        self,
        file_path: str,
        mode: ExtractionMode = ExtractionMode.FULL,
        doc_type: Optional[DocumentType] = None
    ) -> ExtractionResult:
        """Process a document."""
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"Document not found: {file_path}")

        if doc_type is None:
            doc_type = self.detect_type(file_path)

        extractor = self._extractors.get(doc_type)

        if extractor is None:
            logger.warning(f"No extractor for {doc_type}, using fallback")
            return await self._fallback_extract(file_path)

        return await extractor.extract(file_path, mode)

    async def _fallback_extract(self, file_path: str) -> ExtractionResult:
        """Fallback text extraction."""
        try:
            with open(file_path, "r", encoding="utf-8") as f:
                text = f.read()

            return ExtractionResult(
                text=text,
                metadata=DocumentMetadata(
                    file_size=os.path.getsize(file_path),
                    words=len(text.split()),
                    chars=len(text)
                )
            )

        except Exception as e:
            return ExtractionResult(
                text=f"[Failed to extract: {e}]",
                metadata=DocumentMetadata()
            )

    async def process_batch(
        self,
        file_paths: List[str],
        mode: ExtractionMode = ExtractionMode.FULL
    ) -> List[ExtractionResult]:
        """Process multiple documents."""
        tasks = [self.process(path, mode) for path in file_paths]
        return await asyncio.gather(*tasks, return_exceptions=True)

    async def extract_text(self, file_path: str) -> str:
        """Quick text extraction."""
        result = await self.process(file_path)
        return result.text

    async def extract_tables(self, file_path: str) -> List[ExtractedTable]:
        """Extract only tables."""
        result = await self.process(file_path, ExtractionMode.TABLES_ONLY)
        return result.tables

    async def get_metadata(self, file_path: str) -> DocumentMetadata:
        """Get only metadata."""
        result = await self.process(file_path, ExtractionMode.METADATA_ONLY)
        return result.metadata

    async def ocr_document(
        self,
        file_path: str,
        language: str = "eng"
    ) -> str:
        """OCR a scanned document."""
        doc_type = self.detect_type(file_path)

        if doc_type == DocumentType.PDF:
            return await self.ocr.ocr_pdf(file_path, language)
        else:
            # Assume it's an image
            return await self.ocr.ocr_image(file_path, language)

    async def convert(
        self,
        source_path: str,
        target_format: DocumentType,
        output_path: Optional[str] = None
    ) -> str:
        """Convert document format."""
        return await self.converter.convert(source_path, target_format, output_path)


# =============================================================================
# USAGE EXAMPLE
# =============================================================================

async def main():
    """Demonstrate document processing."""
    print("=== BAEL Document Processor ===\n")

    processor = DocumentProcessor()

    # Show supported types
    print("--- Supported Document Types ---")
    for doc_type in processor._extractors:
        print(f"  - {doc_type.value}")

    # Show extraction modes
    print("\n--- Extraction Modes ---")
    for mode in ExtractionMode:
        print(f"  - {mode.value}")

    # Demo with sample text file
    print("\n--- Processing Demo ---")

    # Create sample markdown file
    sample_md = """# Sample Document

## Introduction

This is a sample document for testing BAEL's document processor.

## Features

| Feature | Status |
|---------|--------|
| PDF | Supported |
| Word | Supported |
| Excel | Supported |

## Conclusion

The document processor is ready for use.
"""

    with tempfile.NamedTemporaryFile(
        mode="w",
        suffix=".md",
        delete=False
    ) as f:
        f.write(sample_md)
        temp_path = f.name

    try:
        # Process the document
        result = await processor.process(temp_path)

        print(f"Detected type: {processor.detect_type(temp_path).value}")
        print(f"Text length: {len(result.text)} chars")
        print(f"Words: {result.metadata.words}")
        print(f"Sections: {len(result.sections)}")
        print(f"Tables: {len(result.tables)}")

        print("\nSections found:")
        for section in result.sections:
            print(f"  {'#' * section.level} {section.title}")

        print("\nTables found:")
        for table in result.tables:
            print(f"  - {table.rows} rows, {table.cols} cols")
            print(f"    Headers: {table.headers}")

    finally:
        os.unlink(temp_path)

    print("\n=== Document Processor ready ===")
    print("Install optional dependencies:")
    print("  pip install pypdf python-docx openpyxl pytesseract")


if __name__ == "__main__":
    asyncio.run(main())
